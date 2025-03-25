import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml import OxmlElement, ns
import numpy as np
import seaborn as sns
import os
from collections import defaultdict
from sklearn.metrics import confusion_matrix
import math

def set_cell_background_color(cell, color):
    cell._element.get_or_add_tcPr()
    shading_xml = f'<w:shd {nsdecls("w")} w:fill="{color}"/>'
    cell._element.get_or_add_tcPr().append(parse_xml(shading_xml))

def generate_sbbox_report(csv_path: str, error_log_path: str, IOU_avg : float, Precision_avg : float, Recall_avg : float,
                          F1score_avg : float, F2score_avg : float, choice: str,
                          report_path: str = "Similarity_Report_BBox.docx", threshold: float = 0.5):
    
    df = pd.read_csv(csv_path)
    error_df = pd.read_csv(error_log_path)

    total_pairs = len(df)
    problems_count = len(error_df)
    correct_bounding_boxes = len(df[df['IoU'] >= threshold])
    incorrect_bounding_boxes = len(df[df['IoU'] < threshold])

    if os.path.exists(report_path):
        doc = Document(report_path)
    else:
        doc = Document()
        doc.add_paragraph("Bounding Boxes Detection Report", style="Title").alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    
    doc.add_paragraph("Bounding Boxes Detection Statistics:", style= "Heading 1")
    table = doc.add_table(rows=9, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    stats = [
        ("Total Pairs Count", total_pairs),
        ("Average IOU", round(IOU_avg, 4)),
        ("Average Precision", round(Precision_avg, 4)),
        ("Average Recall", round(Recall_avg, 4)),
        ("Average F1-Score", round(F1score_avg, 4)),
        ("Average F2-Score", round(F2score_avg, 4)),
        ("Count of Correct Predicted Bounding Boxes", correct_bounding_boxes),
        ("Count of Incorrect Predicted Bounding Boxes", incorrect_bounding_boxes),
        ("Count of Problems", problems_count)
    ]

    for i, (label, value) in enumerate(stats):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = str(value)
        set_cell_background_color(table.cell(i, 0), "D3D3D3")

    doc.add_paragraph("Metrics Explained", style="Heading 1")
    paragraph = doc.add_paragraph("Intersection Over Union IOU:", style="Heading 2")
    for run in paragraph.runs:
        run.underline = True
    doc.add_paragraph(
        "The mean IoU (mIoU) gives you a single number that summarizes how well your model's predicted regions "
        "overlap with the ground truth regions across the entire dataset.\n\n"
        "Here’s how to interpret the mIoU score:"
    )
    doc.add_paragraph("1.0 (100%) — Perfect overlap: The predicted regions match exactly with the ground truth for all samples.", style="List Bullet")
    doc.add_paragraph("0.0 (0%) — No overlap at all: The predicted regions and ground truth don’t intersect.", style="List Bullet")
    doc.add_paragraph(
        "In Short:"
    )
    doc.add_paragraph("Higher mIoU → Better overall model performance.", style="List Bullet")
    doc.add_paragraph("Lower mIoU → More room for improvement in predictions.", style="List Bullet")

    #chart
    plt.figure(figsize=(8, 5))
    plt.hist(df['IoU'], bins=10, color='skyblue', edgecolor='black')
    plt.title('IoU Scores Distribution')
    plt.xlabel('IoU Scores')
    plt.ylabel('Frequency')
    plt.axvline(x=0.5, color='red', linestyle='--', label='IoU = 0.5 Threshold')
    plt.legend()

    chart_path = "IOU.png"
    plt.savefig(chart_path)
    plt.close()
    doc.add_picture(chart_path, width=Inches(5))
    doc.add_paragraph('Threshold = 0.5. It’s a balanced choice, doesn’t demand perfect overlap but filters out loose predictions.')

    if choice == '3':
        augmentations = df["Augmentation"].unique()

        for aug in augmentations:
            subset = df[df["Augmentation"] == aug]
            
            count_exceeds = (subset["IoU"] >= threshold).sum()
            count_below = (subset["IoU"] < threshold).sum()
            
            labels = [f'IoU ≥ {threshold}', f'IoU < {threshold}']
            sizes = [count_exceeds, count_below]
            colors = ['#4caf50', '#f44336']
            explode = (0.1, 0)
            plt.figure(figsize=(8, 6))
            plt.pie(sizes, labels=labels, autopct='%1.1f%%', colors=colors, explode=explode, startangle=140)
            plt.title(f'IoU Distribution for {aug}')
            chart_path = "IOU_augmentations.png"
            plt.savefig(chart_path)
            plt.close()
            doc.add_picture(chart_path, width=Inches(5))

    doc.add_paragraph()

    paragraph = doc.add_paragraph("Precision:", style="Heading 2")
    for run in paragraph.runs:
        run.underline = True
    doc.add_paragraph("High: Most detected bboxes are correct (few false positives).", style = "List Bullet")
    doc.add_paragraph("Low: Many detected bboxes are incorrect (many false positives).", style = "List Bullet")

    doc.add_paragraph()

    paragraph = doc.add_paragraph("Recall:", style="Heading 2")
    for run in paragraph.runs:
        run.underline = True
    doc.add_paragraph("High: Most true faces are detected (few false negatives).", style = "List Bullet")
    doc.add_paragraph("Low: Many faces are missed (many false negatives).", style = "List Bullet")

    doc.add_paragraph()

    paragraph = doc.add_paragraph("F1 Score:", style="Heading 2")
    for run in paragraph.runs:
        run.underline = True
    doc.add_paragraph("High: Good balance between catching all faces and ensuring detections are accurate.", style = "List Bullet")
    doc.add_paragraph("Low: Either too many false positives or too many missed faces.", style = "List Bullet")

    doc.add_paragraph()

    paragraph = doc.add_paragraph("F2 Score:", style="Heading 2")
    for run in paragraph.runs:
        run.underline = True
    doc.add_paragraph("Like F1, but weighs recall higher.", style = "List Bullet")
    doc.add_paragraph("High: Strong focus on minimizing missed detections.", style = "List Bullet")
    doc.add_paragraph("Low: Many faces are missed.", style = "List Bullet")

    doc.add_paragraph()
    
    
    # Error Log
    doc.add_paragraph("Error Log:", style="Heading 1")

    # to collapse or expand the section
    paragraph = doc.paragraphs[-1]
    run = paragraph.runs[0]
    rPr = run._element.get_or_add_rPr()
    collapse = OxmlElement('w:collapsed')
    collapse.set(ns.qn('w:val'), '1')
    rPr.append(collapse)

    error_table = doc.add_table(rows=1, cols=2)
    error_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    error_table.style = 'Table Grid'  
    tbl = error_table._element
    tblPr = tbl.find(ns.qn('w:tblPr'))  
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.append(tblPr)

    tblW = OxmlElement('w:tblW')
    tblW.set(ns.qn('w:w'), "5000")  
    tblW.set(ns.qn('w:type'), "pct")  
    tblPr.append(tblW)
    column_widths = [Inches(3), Inches(2.5)]
    hdr_cells = error_table.rows[0].cells
    headers = ["Error Message", "Image Path"]
    for i, text in enumerate(headers):
        hdr_cells[i].text = text
        hdr_cells[i].paragraphs[0].runs[0].bold = True  
        hdr_cells[i].paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER  
        if i == 0: 
            tcPr = hdr_cells[i]._element.find(ns.qn('w:tcPr'))
            if tcPr is None:
                tcPr = OxmlElement('w:tcPr')
                hdr_cells[i]._element.append(tcPr)

            noWrap = OxmlElement('w:noWrap')
            tcPr.append(noWrap)
    unique_errors = set()
    for _, row in error_df.iterrows():
        key = (row["Error"], row["Image Path"])
        if key not in unique_errors:
            unique_errors.add(key)
            error_row = error_table.add_row().cells
            error_row[0].text = row["Error"]
            error_row[1].text = row["Image Path"]

        for i, cell in enumerate(error_row):
            cell.width = column_widths[i]
            cell.paragraphs[0].runs[0].font.size = Pt(10) 
            cell.paragraphs[0].alignment = WD_TABLE_ALIGNMENT.CENTER

    doc.save(report_path)
    print(f"Face Verification Model Report generated at {report_path}")
    return report_path