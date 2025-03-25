import pandas as pd
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches

def generate_detection_report(csv_path, report_path="Detection_Metrics_Report.docx"):
    """Generate a report for detection evaluation metrics from the dataset."""
    # Load CSV dataset
    df = pd.read_csv(csv_path)
 
    # Ensure required columns exist
    required_columns = ["IoU", "Precision", "Recall", "F1 Score", "F2 Score", "F1 Score Width", "F2 Score Width"]
    if not all(col in df.columns for col in required_columns):
        raise ValueError(f"Missing required columns in CSV. Expected: {required_columns}")
 
    # Create Document
    doc = Document()
    doc.add_paragraph("Object Detection Metrics Report", style="Title").alignment = 1
 
    # 📌 **Dataset Overview**
    total_data_points = len(df)
    missing_data = df.isnull().sum().sum()
    doc.add_paragraph("Dataset Overview", style="Heading 1")
    doc.add_paragraph(f"Total Data Points: {total_data_points}")
    doc.add_paragraph(f"Missing Values: {missing_data}")
 
    # 📊 **Evaluation Metrics Summary**
    doc.add_paragraph("Evaluation Metrics Summary", style="Heading 1")
    doc.add_paragraph("This table summarizes the object detection performance metrics across the dataset.")
 
    # Create Summary Table
    metrics_summary = df[required_columns].mean().to_frame().T
    table = doc.add_table(rows=2, cols=len(required_columns) + 1)
    table.alignment = 1
 
    # Table Headers
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Metric"
    for j, metric in enumerate(required_columns):
        hdr_cells[j + 1].text = metric
 
    # Table Data
    row_cells = table.rows[1].cells
    row_cells[0].text = "Mean Score"
    for j, metric in enumerate(required_columns):
        row_cells[j + 1].text = f"{metrics_summary[metric].values[0]:.4f}"
 
    # 📊 **Generate a bar chart for metrics**
    metrics_summary.plot(kind='bar', figsize=(10, 5), colormap="Set2")
    plt.title("Average Detection Metrics")
    plt.ylabel("Score")
    plt.xticks([])
    plt.grid(axis='y', linestyle="--", alpha=0.7)
    metrics_chart_path = "detection_metrics_chart.png"
    plt.savefig(metrics_chart_path)
    plt.close()
 
    doc.add_paragraph("\nThe bar chart below shows a summary of object detection performance metrics:")
    doc.add_picture(metrics_chart_path, width=Inches(5))
 
    # 📌 **Error Analysis**
    doc.add_paragraph("Error Analysis", style="Heading 1")
 
    # Compute different error categories
    perfect_detections = df[df["IoU"] >= 0.9]  # Perfect matches
    failed_detections = df[df["IoU"] < 0.5]  # Poor matches
    high_error = df[df["IoU"] < 0.3]  # High error detections
 
    total = len(df)
    error_counts = {
        "Perfect Detections (IoU ≥ 0.9)": (len(perfect_detections), len(perfect_detections) / total * 100),
        "Failed Detections (IoU < 0.5)": (len(failed_detections), len(failed_detections) / total * 100),
        "High Error (IoU < 0.3)": (len(high_error), len(high_error) / total * 100),
    }
 
    # **Create Error Analysis Table**
    error_table = doc.add_table(rows=len(error_counts) + 1, cols=3)
    error_table.alignment = 1
 
    # Table Headers
    error_hdr_cells = error_table.rows[0].cells
    error_hdr_cells[0].text = "Category"
    error_hdr_cells[1].text = "Count"
    error_hdr_cells[2].text = "Percentage"
 
    # Table Data
    for i, (category, (count, percentage)) in enumerate(error_counts.items()):
        row_cells = error_table.rows[i + 1].cells
        row_cells[0].text = category
        row_cells[1].text = str(count)
        row_cells[2].text = f"{percentage:.2f}%"
 
    # 📊 **Generate Error Analysis Chart**
    error_analysis_df = pd.DataFrame(error_counts, index=["Count", "Percentage"]).T[["Count"]]
    error_analysis_df.plot(kind='bar', figsize=(10, 5), colormap="coolwarm")
    plt.title("Error Analysis: Detection Performance")
    plt.ylabel("Count")
    plt.xticks(rotation=45)
    plt.grid(axis='y', linestyle="--", alpha=0.7)
    error_chart_path = "detection_error_chart.png"
    plt.savefig(error_chart_path)
    plt.close()
 
    doc.add_paragraph("\nThe bar chart below illustrates detection errors in the dataset:")
    doc.add_picture(error_chart_path, width=Inches(5))
 
    # **Save the final report**
    doc.save(report_path)
    print(f"✅ Detection Metrics Report generated at {report_path}")
    return report_path